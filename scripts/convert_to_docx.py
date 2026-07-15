import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Colors
ORANGE_HEX = "F97316"
ORANGE_RGB = RGBColor(249, 115, 22)
DARK_HEX = "111111"
DARK_RGB = RGBColor(17, 17, 17)
GRAY_HEX = "666666"
GRAY_RGB = RGBColor(102, 102, 102)
LIGHT_GRAY_HEX = "F5F5F5"
BORDER_HEX = "E5E5E5"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    # 20 dxa = 1 pt. Default margins: top/bottom ~ 7pt, left/right ~ 9pt
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_paragraph_left_border(paragraph, hex_color="F97316", sz="24"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)  # 24 = 3pt
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)

def add_formatted_text(paragraph, text, default_font="Segoe UI", default_size=11, is_code_block=False):
    # Tokenize bold (**), italic (*), code (`), and links ([text](url))
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
        run = paragraph.add_run()
        run.font.name = default_font
        run.font.size = Pt(default_size)
        
        if token.startswith('**') and token.endswith('**'):
            run.text = token[2:-2]
            run.bold = True
            run.font.color.rgb = DARK_RGB
        elif token.startswith('*') and token.endswith('*'):
            run.text = token[1:-1]
            run.italic = True
            run.font.color.rgb = DARK_RGB
        elif token.startswith('`') and token.endswith('`'):
            run.text = token[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(default_size - 1)
            run.font.color.rgb = RGBColor(199, 37, 78) # standard code red
        elif token.startswith('[') and ']' in token and token.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if match:
                run.text = match.group(1)
                run.font.color.rgb = ORANGE_RGB
                run.underline = True
        else:
            run.text = token
            run.font.color.rgb = DARK_RGB

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), BORDER_HEX)
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def build_docx(md_path, docx_path):
    doc = Document()
    
    # Configure document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. CREATE TITLE PAGE (COVER)
    for _ in range(3):
        doc.add_paragraph() # space

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BÁO CÁO PHÂN TÍCH DỰ ÁN BLOACT")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(28)
    title_run.bold = True
    title_run.font.color.rgb = ORANGE_RGB

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("NỀN TẢNG BLOG CÁ NHÂN TỐI GIẢN & CHUẨN SEO\n(Next.js App Router & Supabase)")
    sub_run.font.name = 'Segoe UI'
    sub_run.font.size = Pt(14)
    sub_run.italic = True
    sub_run.font.color.rgb = GRAY_RGB

    for _ in range(8):
        doc.add_paragraph() # space

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        "Tài liệu hướng dẫn bảo vệ Đồ án tốt nghiệp / Bài tập lớn\n"
        "Vai trò: Senior System Architect & Hướng dẫn kỹ thuật\n"
        "Ngôn ngữ phân tích: Tiếng Việt"
    )
    info_run.font.name = 'Segoe UI'
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = GRAY_RGB

    doc.add_page_break()

    # 2. READ AND PARSE THE MARKDOWN FILE
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_text = []
    code_lang = ""
    
    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Handle Code Block
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block, render it
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(12)
                
                # Combine lines and add as monospace
                full_code = "\n".join(code_text)
                
                # Set background shading for code block
                set_paragraph_left_border(p, hex_color="666666", sz="12")
                
                run = p.add_run(full_code)
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK_RGB
                
                # Add extra spacing after code block
                code_text = []
            else:
                # Start of code block
                in_code_block = True
                code_lang = stripped[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_text.append(line.rstrip('\n'))
            i += 1
            continue

        # Handle Table
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table, render it
            in_table = False
            # Parse table rows
            parsed_rows = []
            for tr in table_rows:
                # split by | and clean up
                cols = [c.strip() for c in tr.split('|')[1:-1]]
                parsed_rows.append(cols)
            
            table_rows = []
            
            if len(parsed_rows) > 0:
                # Remove separator row (e.g. |---|---|)
                if len(parsed_rows) > 1 and all(all(ch in '-:' for ch in c) for c in parsed_rows[1] if c):
                    header = parsed_rows[0]
                    body = parsed_rows[2:]
                else:
                    header = parsed_rows[0]
                    body = parsed_rows[1:]
                
                cols_count = len(header)
                table = doc.add_table(rows=0, cols=cols_count)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                # Render Header
                hdr_row = table.add_row()
                for c_idx, cell_text in enumerate(header):
                    cell = hdr_row.cells[c_idx]
                    set_cell_background(cell, ORANGE_HEX)
                    set_cell_margins(cell)
                    cell_p = cell.paragraphs[0]
                    cell_p.paragraph_format.space_before = Pt(2)
                    cell_p.paragraph_format.space_after = Pt(2)
                    run = cell_p.add_run(cell_text)
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10)
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                
                # Render Body
                for r_idx, row_data in enumerate(body):
                    if len(row_data) < cols_count:
                        row_data += [""] * (cols_count - len(row_data))
                    row_data = row_data[:cols_count]
                    
                    row = table.add_row()
                    bg_color = LIGHT_GRAY_HEX if r_idx % 2 == 1 else "FFFFFF"
                    for c_idx, cell_text in enumerate(row_data):
                        cell = row.cells[c_idx]
                        if bg_color != "FFFFFF":
                            set_cell_background(cell, bg_color)
                        set_cell_margins(cell)
                        cell_p = cell.paragraphs[0]
                        cell_p.paragraph_format.space_before = Pt(2)
                        cell_p.paragraph_format.space_after = Pt(2)
                        add_formatted_text(cell_p, cell_text, default_size=9.5)
                
                # Add spacer after table
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(12)
            
            # Note: do not increment i here, as this line might be a regular line
            continue

        # Skip separator line (---)
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Add a bottom border XML or just an actual horizontal line
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), BORDER_HEX)
            pBdr.append(bottom)
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Handle Images
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            caption = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve relative paths
            if not os.path.isabs(img_path):
                img_path = os.path.join(os.path.dirname(md_path), img_path)
            
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                
                # Add picture
                try:
                    p.add_run().add_picture(img_path, width=Inches(5.5))
                except Exception as e:
                    run = p.add_run(f"[Lỗi hiển thị ảnh: {str(e)}]")
                    run.font.color.rgb = RGBColor(239, 68, 68)
                
                # Add caption
                if caption:
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_before = Pt(2)
                    cap_p.paragraph_format.space_after = Pt(12)
                    cap_run = cap_p.add_run(f"Hình: {caption}")
                    cap_run.font.name = 'Segoe UI'
                    cap_run.font.size = Pt(9.5)
                    cap_run.italic = True
                    cap_run.font.color.rgb = GRAY_RGB
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Hình ảnh không tìm thấy: {img_path}]")
                run.font.color.rgb = RGBColor(239, 68, 68)
            i += 1
            continue

        # Headers
        if stripped.startswith("#"):
            h_level = len(re.match(r'^#+', stripped).group(0))
            h_text = stripped[h_level:].strip()
            
            # Clean up numberings if duplicate
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if h_level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(h_text)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(18)
                run.bold = True
                run.font.color.rgb = ORANGE_RGB
            elif h_level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(h_text)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(14)
                run.bold = True
                run.font.color.rgb = ORANGE_RGB
            else:
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(h_text)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = DARK_RGB
            i += 1
            continue

        # Blockquote (> text)
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_left_border(p, hex_color=ORANGE_HEX, sz="24")
            
            quote_text = stripped[1:].strip()
            # If quote text starts with alert prefixes like [!NOTE], [!IMPORTANT], etc.
            alert_match = re.match(r'^\[!(NOTE|IMPORTANT|WARNING|TIP|CAUTION)\]', quote_text)
            if alert_match:
                alert_type = alert_match.group(1)
                quote_text = quote_text[len(alert_match.group(0)):].strip()
                run_alert = p.add_run(f"[{alert_type}] ")
                run_alert.bold = True
                run_alert.font.name = 'Segoe UI'
                run_alert.font.size = Pt(10.5)
                if alert_type in ['IMPORTANT', 'WARNING', 'CAUTION']:
                    run_alert.font.color.rgb = RGBColor(239, 68, 68) # Red
                else:
                    run_alert.font.color.rgb = ORANGE_RGB
            
            start_idx = len(p.runs)
            add_formatted_text(p, quote_text, default_size=10.5)
            for r in p.runs[start_idx:]:
                r.italic = True
                if r.font.color.rgb is None or r.font.color.rgb == DARK_RGB:
                    r.font.color.rgb = GRAY_RGB
            i += 1
            continue

        # Lists (unordered - or * or numbered 1.)
        list_match = re.match(r'^(\*|-|\+)\s+(.*)', stripped)
        num_list_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        
        if list_match:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, list_match.group(2))
            i += 1
            continue
        elif num_list_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            num_run = p.add_run(f"{num_list_match.group(1)}. ")
            num_run.bold = True
            num_run.font.name = 'Segoe UI'
            num_run.font.size = Pt(11)
            num_run.font.color.rgb = DARK_RGB
            
            add_formatted_text(p, num_list_match.group(2))
            i += 1
            continue

        # Normal Paragraph
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, stripped)
        
        i += 1

    doc.save(docx_path)
    print(f"Successfully generated {docx_path}")

if __name__ == "__main__":
    brain_dir = r"C:\Users\Administrator\.gemini\antigravity\brain\d27d6486-b4ba-4850-b883-f0055531e98d"
    md_file = os.path.join(brain_dir, "artifacts", "bao_cao_phan_tich_bloact.md")
    docx_file = os.path.join(brain_dir, "artifacts", "bao_cao_phan_tich_bloact.docx")
    
    # Save a copy in workspace too for easy access
    workspace_docx_file = r"d:\Bloact\bao_cao_phan_tich_bloact.docx"
    
    if os.path.exists(md_file):
        build_docx(md_file, docx_file)
        build_docx(md_file, workspace_docx_file)
    else:
        print("Markdown file not found!")
